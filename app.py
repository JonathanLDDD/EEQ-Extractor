


import streamlit as st
import os
import fitz  # PyMuPDF for PDF
from docx import Document
from openai import OpenAI

# ========== GPT Client ==========
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
MODEL_NAME = "gpt-4-1106-preview"

# ========== 工具函数 ==========
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_pdf(file_path):
    text = ""
    pdf = fitz.open(file_path)
    for page in pdf:
        text += page.get_text()
    return text

def ask_gpt(prompt):
    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[
            {"role": "system", "content": "You are an assistant helping QA Commons extract and format EEQs."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.2
    )
    return response.choices[0].message.content

def write_output_to_word(output_text, output_path):
    doc = Document()
    doc.add_heading("EEQ Extracted Output", 0)
    doc.add_paragraph(output_text)
    doc.save(output_path)

# ========== Streamlit App ==========
st.title("EEQ Extractor")
uploaded_file = st.file_uploader("Upload syllabus (PDF or Word)", type=["pdf", "docx"])

if uploaded_file is not None:
    file_path = "temp_" + os.path.basename(uploaded_file.name)
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    if file_path.endswith(".pdf"):
        syllabus_text = extract_text_from_pdf(file_path)
    else:
        syllabus_text = extract_text_from_docx(file_path)

    with open("Prompt.txt", "r", encoding="utf-8") as f:
        base_prompt = f.read()
    full_prompt = base_prompt + "\n\n" + syllabus_text

    with st.spinner("Processing with GPT..."):
        gpt_output = ask_gpt(full_prompt)

    output_file = "EEQ_output.docx"
    write_output_to_word(gpt_output, output_file)

    with open(output_file, "rb") as f:
        st.download_button("Download Processed EEQ", f, file_name="EEQ_output.docx")

    st.success("Processing complete!")
